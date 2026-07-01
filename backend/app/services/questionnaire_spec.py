"""Questionnaire programming spec export (Excel / Word) for field programmers."""

from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.shared import Inches
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


def _spec_rows(schema: dict[str, Any]) -> list[dict[str, str]]:
    groups = {g["id"]: g.get("title", "") for g in schema.get("groups") or []}
    rows: list[dict[str, str]] = []
    for var in schema.get("variables") or []:
        gid = var.get("group_id")
        options = var.get("answer_options") or []
        option_text = "; ".join(
            f"{o.get('code', '')}={o.get('label', '')}" for o in options[:40]
        )
        subs = var.get("subquestions") or []
        sub_text = "; ".join(f"{s.get('code', '')}={s.get('label', '')}" for s in subs[:20])
        rows.append(
            {
                "group": str(groups.get(gid) or var.get("group_title") or ""),
                "code": str(var.get("code") or ""),
                "variable_id": str(var.get("id") or ""),
                "question": str(var.get("text") or ""),
                "type": str(var.get("type_label") or var.get("ls_type") or ""),
                "kind": str(var.get("kind") or ""),
                "answers": option_text or sub_text,
                "columns": ", ".join(str(c) for c in (var.get("columns") or [])[:8]),
            }
        )
    return rows


def questionnaire_spec_excel(schema: dict[str, Any], *, title: str = "Questionnaire") -> bytes:
    rows = _spec_rows(schema)
    wb = Workbook()
    ws = wb.active
    ws.title = "Questionnaire"
    headers = ["Group", "Code", "Variable ID", "Question text", "Type", "Kind", "Answers / items", "Columns"]

    ws.merge_cells("A1:H1")
    title_cell = ws.cell(1, 1, title)
    title_cell.font = Font(bold=True, size=14, color="0B2545", name="Calibri")
    title_cell.alignment = Alignment(horizontal="left", vertical="center")

    header_fill = PatternFill("solid", fgColor="0B2545")
    header_font = Font(bold=True, color="FFFFFF", name="Calibri")
    thin = Side(style="thin", color="CBD5E1")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    zebra = PatternFill("solid", fgColor="F8FAFC")

    header_row = 2
    for col, label in enumerate(headers, 1):
        cell = ws.cell(row=header_row, column=col, value=label)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
        cell.border = border

    for r, row in enumerate(rows, header_row + 1):
        values = [
            row["group"],
            row["code"],
            row["variable_id"],
            row["question"],
            row["type"],
            row["kind"],
            row["answers"],
            row["columns"],
        ]
        fill = zebra if (r - header_row) % 2 == 0 else None
        for col, value in enumerate(values, 1):
            cell = ws.cell(row=r, column=col, value=value)
            cell.font = Font(size=10, name="Calibri")
            cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
            cell.border = border
            if fill:
                cell.fill = fill

    ws.freeze_panes = "A3"
    ws.sheet_view.showGridLines = False
    widths = [22, 12, 14, 48, 16, 12, 40, 24]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    meta = wb.create_sheet("Meta")
    meta["A1"] = "Survey"
    meta["B1"] = title
    meta["A2"] = "Questions"
    meta["B2"] = len(rows)
    meta.column_dimensions["A"].width = 14
    meta.column_dimensions["B"].width = 40

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def questionnaire_spec_docx(schema: dict[str, Any], *, title: str = "Questionnaire") -> bytes:
    rows = _spec_rows(schema)
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(f"{len(rows)} questions — programming specification for LimeSurvey / field team.")

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, label in enumerate(["Code", "Type", "Question", "Answers / items", "Group"]):
        hdr[i].text = label
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.bold = True

    for row in rows:
        cells = table.add_row().cells
        cells[0].text = row["code"]
        cells[1].text = row["type"]
        cells[2].text = row["question"][:500]
        cells[3].text = row["answers"][:800]
        cells[4].text = row["group"]

    for section in doc.sections:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
